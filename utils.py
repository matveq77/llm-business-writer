import os
import json
import google.generativeai as genai
from docx import Document
from pydantic import BaseModel
from dotenv import load_dotenv

load_dotenv()

api_key = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=api_key)

MODEL_NAME = "models/gemini-1.5-flash" 

model = genai.GenerativeModel(model_name=MODEL_NAME)

def get_order_data_gemini(user_text: str) -> dict:

    prompt = f"""
    Извлеки данные для приказа о командировке из текста: "{user_text}"
    Верни ответ строго в формате JSON с полями:
    employee_name, destination, duration_days, start_date, reason.
    Если какая-то информация отсутствует, напиши "Не указано".
    Используй русский язык для значений.
    """
    
    try:
        response = model.generate_content(
            prompt,
            generation_config={
                "response_mime_type": "application/json",
                "temperature": 0.1
            }
        )
        return json.loads(response.text)
    except Exception as e:
        print(f"Ошибка внутри Gemini: {e}")
        print("Список доступных вам моделей:")
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                print(f"- {m.name}")
        raise e

def create_docx(data: dict):
    doc = Document()
    doc.add_heading('ПРИКАЗ О КОМАНДИРОВКЕ', 0)
    
    table = doc.add_table(rows=5, cols=2)
    rows = [
        ("Сотрудник:", str(data.get("employee_name", "Не указано"))),
        ("Место назначения:", str(data.get("destination", "Не указано"))),
        ("Срок (дней):", str(data.get("duration_days", "Не указано"))),
        ("Дата начала:", str(data.get("start_date", "Не указано"))),
        ("Цель поездки:", str(data.get("reason", "Не указано"))),
    ]
    
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    file_name = f"order_{data.get('employee_name', 'document')}.docx".replace(" ", "_")
    file_path = os.path.join(os.getcwd(), file_name)
    doc.save(file_path)
    return file_path